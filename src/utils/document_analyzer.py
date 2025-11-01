"""
문서 분석 및 체크리스트 변환 모듈

PDF/Word 문서를 읽어서 주요 의사결정 항목을 추출하고
체크리스트 항목으로 변환합니다.
"""

import os
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging

# PDF 처리 라이브러리
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2가 설치되지 않았습니다. PDF 처리가 제한됩니다.")

# Word 문서 처리 라이브러리
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx가 설치되지 않았습니다. Word 문서 처리가 제한됩니다.")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DocumentAnalyzer:
    """문서 분석기 클래스"""

    def __init__(self):
        """초기화"""
        self.keywords = {
            '승인': ['승인', '인허가', '허가', '면허', '등록', '신고', '협의'],
            '비용': ['비용', 'CAPEX', 'OPEX', '예산', '투자', '지출', '경비'],
            '일정': ['일정', '공정', '기한', '납기', '완료', '착수', '준공'],
            '환경': ['환경', 'EIA', '환경영향평가', '소음', '대기', '수질', '폐기물'],
            '안전': ['안전', '위험', '사고', '재해', '보안', '화재', '방재'],
            '운영': ['운영', 'OTP', '수하물', '회전율', '용량', '처리량', '서비스'],
            '설계': ['설계', '구조', '배치', '레이아웃', '시설', '장비', '시스템'],
            '계획': ['계획', '전략', '방침', '정책', '기본계획', '실행계획']
        }

    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """
        문서를 분석하여 체크리스트 항목을 추출

        Args:
            file_path: 분석할 문서 파일 경로

        Returns:
            Dict: 추출된 정보
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        # 파일 확장자에 따라 적절한 메서드 호출
        ext = file_path.suffix.lower()

        if ext == '.pdf':
            text = self._extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            text = self._extract_text_from_docx(file_path)
        elif ext == '.txt':
            text = self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")

        # 텍스트 분석
        analysis_result = self._analyze_text(text)

        return {
            'file_name': file_path.name,
            'file_path': str(file_path),
            'text': text,
            'analysis': analysis_result
        }

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """PDF에서 텍스트 추출"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2가 설치되지 않았습니다. 'pip install PyPDF2'로 설치하세요.")

        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"PDF 텍스트 추출 실패: {e}")
            raise

        return text

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Word 문서에서 텍스트 추출"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx가 설치되지 않았습니다. 'pip install python-docx'로 설치하세요.")

        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # 표 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"Word 문서 텍스트 추출 실패: {e}")
            raise

        return text

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """텍스트 파일에서 내용 추출"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # UTF-8 실패 시 다른 인코딩 시도
            with open(file_path, 'r', encoding='cp949') as file:
                text = file.read()

        return text

    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """
        텍스트를 분석하여 주요 항목 추출

        Args:
            text: 분석할 텍스트

        Returns:
            Dict: 분석 결과
        """
        # 문장 단위로 분리
        sentences = self._split_into_sentences(text)

        # 키워드 기반 분류
        categorized_sentences = self._categorize_sentences(sentences)

        # 체크리스트 항목 후보 추출
        checklist_candidates = self._extract_checklist_candidates(categorized_sentences)

        return {
            'total_sentences': len(sentences),
            'categorized_sentences': categorized_sentences,
            'checklist_candidates': checklist_candidates
        }

    def _split_into_sentences(self, text: str) -> List[str]:
        """텍스트를 문장 단위로 분리"""
        # 줄바꿈 정리
        text = re.sub(r'\n+', '\n', text)

        # 문장 분리 (한글 문장 부호 고려)
        sentences = re.split(r'[.!?]\s+|\n', text)

        # 빈 문장 제거 및 정리
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]

        return sentences

    def _categorize_sentences(self, sentences: List[str]) -> Dict[str, List[str]]:
        """문장을 카테고리별로 분류"""
        categorized = {category: [] for category in self.keywords.keys()}
        categorized['기타'] = []

        for sentence in sentences:
            matched = False
            for category, keywords in self.keywords.items():
                if any(keyword in sentence for keyword in keywords):
                    categorized[category].append(sentence)
                    matched = True
                    break

            if not matched:
                categorized['기타'].append(sentence)

        # 빈 카테고리 제거
        categorized = {k: v for k, v in categorized.items() if v}

        return categorized

    def _extract_checklist_candidates(self, categorized_sentences: Dict[str, List[str]]) -> List[Dict[str, Any]]:
        """체크리스트 항목 후보 추출"""
        candidates = []
        item_id = 1

        for category, sentences in categorized_sentences.items():
            for sentence in sentences:
                # 의사결정이나 행동이 필요한 문장 추출
                if self._is_actionable(sentence):
                    candidates.append({
                        'id': item_id,
                        'category': category,
                        'item': self._extract_action_item(sentence),
                        'source_text': sentence
                    })
                    item_id += 1

        return candidates

    def _is_actionable(self, sentence: str) -> bool:
        """문장이 실행 가능한 항목인지 판단"""
        # 행동을 나타내는 키워드
        action_keywords = [
            '필요', '수행', '검토', '확인', '승인', '취득', '획득',
            '분석', '평가', '설계', '계획', '실시', '진행', '완료',
            '준비', '마련', '수립', '작성', '제출', '신청', '협의',
            '조사', '측정', '점검', '관리', '운영', '유지', '보수'
        ]

        return any(keyword in sentence for keyword in action_keywords)

    def _extract_action_item(self, sentence: str) -> str:
        """문장에서 실행 항목을 추출하여 체크리스트 형식으로 변환"""
        # 불필요한 접속사나 부사 제거
        sentence = re.sub(r'^(그러므로|따라서|또한|또|그리고|하지만|그러나|즉)\s*', '', sentence)

        # 문장을 간결하게 정리
        if len(sentence) > 100:
            # 너무 긴 문장은 핵심만 추출
            sentence = sentence[:100] + '...'

        return sentence


def load_system_prompt(prompt_path: str = None) -> str:
    """
    시스템 프롬프트 로드

    Args:
        prompt_path: 프롬프트 파일 경로 (None이면 기본 경로 사용)

    Returns:
        str: 시스템 프롬프트 내용
    """
    if prompt_path is None:
        # 기본 경로
        current_dir = Path(__file__).parent.parent.parent
        prompt_path = current_dir / 'config' / 'checklist_system_prompt.md'

    prompt_path = Path(prompt_path)

    if not prompt_path.exists():
        raise FileNotFoundError(f"시스템 프롬프트 파일을 찾을 수 없습니다: {prompt_path}")

    with open(prompt_path, 'r', encoding='utf-8') as f:
        return f.read()


def main():
    """테스트용 메인 함수"""
    analyzer = DocumentAnalyzer()

    # 테스트 문서 경로
    test_file = "test_document.pdf"

    if os.path.exists(test_file):
        result = analyzer.analyze_document(test_file)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"테스트 파일이 없습니다: {test_file}")


if __name__ == "__main__":
    main()
