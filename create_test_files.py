from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Pt

# PDF 파일 생성
def create_test_pdf():
    pdf_filename = "test_risk_management.pdf"
    doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # 제목
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
    )
    title = Paragraph("소프트웨어 개발 프로젝트 위험 관리 계획서", title_style)
    story.append(title)
    story.append(Spacer(1, 0.2*inch))

    # 본문 내용
    content = [
        ("1. 프로젝트 개요", [
            "본 프로젝트는 고객 관리 시스템(CRM) 개발을 목표로 하며, 6개월의 개발 기간과 5명의 개발팀으로 구성됩니다.",
            "주요 기능으로는 고객 데이터 관리, 영업 파이프라인 추적, 보고서 생성 기능이 포함됩니다."
        ]),
        ("2. 기술적 위험 요소", [
            "- 데이터베이스 성능 저하: 대용량 데이터 처리 시 응답 속도가 느려질 수 있습니다.",
            "- API 통합 문제: 외부 결제 시스템과의 연동 과정에서 호환성 문제가 발생할 수 있습니다.",
            "- 보안 취약점: 사용자 인증 및 데이터 암호화 구현 시 보안 결함이 발견될 수 있습니다.",
            "완화 방안: 정기적인 성능 테스트 실시, API 문서 사전 검토, 보안 감사 수행"
        ]),
        ("3. 일정 관련 위험", [
            "- 요구사항 변경: 고객의 추가 요구사항으로 인한 일정 지연 가능성이 있습니다.",
            "- 리소스 부족: 핵심 개발자의 휴가 또는 이직으로 인한 인력 공백이 발생할 수 있습니다.",
            "완화 방안: 변경 관리 프로세스 수립, 백업 인력 확보, 지식 공유 문서화"
        ]),
        ("4. 비용 관련 위험", [
            "- 예산 초과: 추가 기능 개발 및 인프라 비용 증가로 예산을 초과할 수 있습니다.",
            "- 라이선스 비용: 제3자 라이브러리 사용에 따른 예상치 못한 라이선스 비용이 발생할 수 있습니다.",
            "완화 방안: 월별 예산 모니터링, 오픈소스 대안 검토, 예비 예산 확보"
        ]),
        ("5. 품질 관련 위험", [
            "- 테스트 부족: 출시 일정에 쫓겨 충분한 테스트를 수행하지 못할 수 있습니다.",
            "- 기술 부채: 빠른 개발을 위해 임시 해결책을 사용하여 장기적으로 유지보수가 어려워질 수 있습니다.",
            "완화 방안: 자동화 테스트 도입, 코드 리뷰 의무화, 리팩토링 시간 확보"
        ]),
        ("6. 운영 및 유지보수 위험", [
            "- 시스템 다운타임: 서버 장애로 인한 서비스 중단이 발생할 수 있습니다.",
            "- 데이터 손실: 백업 실패 또는 데이터베이스 오류로 고객 데이터가 손실될 수 있습니다.",
            "완화 방안: 이중화 시스템 구축, 자동 백업 설정, 모니터링 알림 시스템 도입"
        ]),
        ("7. 위험 모니터링 계획", [
            "- 주간 위험 검토 회의를 진행하여 새로운 위험을 식별하고 기존 위험의 상태를 점검합니다.",
            "- 위험 관리 대시보드를 통해 실시간으로 주요 지표를 모니터링합니다.",
            "- 월간 보고서를 작성하여 경영진에게 위험 현황을 보고합니다."
        ])
    ]

    for section_title, items in content:
        # 섹션 제목
        section = Paragraph(section_title, styles['Heading2'])
        story.append(section)
        story.append(Spacer(1, 0.1*inch))

        # 섹션 내용
        for item in items:
            para = Paragraph(item, styles['BodyText'])
            story.append(para)
            story.append(Spacer(1, 0.05*inch))

        story.append(Spacer(1, 0.2*inch))

    doc.build(story)
    print(f"PDF 파일 생성 완료: {pdf_filename}")

# DOCX 파일 생성
def create_test_docx():
    doc = Document()

    # 제목
    title = doc.add_heading('클라우드 마이그레이션 프로젝트 위험 관리 체크리스트', 0)

    # 내용
    sections = [
        ("1. 기술 인프라 위험", [
            "□ 클라우드 서비스 호환성 검증 완료",
            "□ 네트워크 대역폭 충분성 확인",
            "□ 데이터 마이그레이션 전략 수립",
            "□ 레거시 시스템 연동 테스트 완료",
            "□ 재해 복구(DR) 계획 수립",
            "",
            "주요 위험:",
            "- 기존 온프레미스 시스템과 클라우드 간 데이터 동기화 실패",
            "- 마이그레이션 중 서비스 중단 시간 초과",
            "- 클라우드 제공업체의 서비스 장애"
        ]),
        ("2. 보안 및 컴플라이언스 위험", [
            "□ 데이터 암호화 정책 수립",
            "□ 접근 권한 관리 체계 구축",
            "□ 규제 준수 요구사항 검토 (GDPR, 개인정보보호법 등)",
            "□ 보안 감사 및 취약점 스캔 실시",
            "□ 백업 및 복구 절차 테스트",
            "",
            "주요 위험:",
            "- 데이터 유출 또는 무단 접근",
            "- 규제 위반으로 인한 법적 제재",
            "- 클라우드 환경의 보안 설정 오류"
        ]),
        ("3. 비용 및 예산 위험", [
            "□ 클라우드 사용 비용 산정 완료",
            "□ 예상 비용과 실제 비용 모니터링 체계 구축",
            "□ 자동 스케일링 정책 설정",
            "□ 미사용 리소스 정리 프로세스 수립",
            "□ 비용 최적화 도구 도입",
            "",
            "주요 위험:",
            "- 예상보다 높은 클라우드 사용 비용 발생",
            "- 자동 스케일링으로 인한 비용 급증",
            "- 숨겨진 비용 항목 발견"
        ]),
        ("4. 운영 및 관리 위험", [
            "□ 클라우드 관리 도구 교육 완료",
            "□ 모니터링 및 알림 시스템 구축",
            "□ 운영 매뉴얼 작성",
            "□ 장애 대응 프로세스 수립",
            "□ 성능 벤치마크 테스트 실시",
            "",
            "주요 위험:",
            "- 운영팀의 클라우드 기술 역량 부족",
            "- 모니터링 공백으로 인한 장애 대응 지연",
            "- 성능 저하 또는 가용성 문제"
        ]),
        ("5. 조직 및 프로세스 위험", [
            "□ 이해관계자 커뮤니케이션 계획 수립",
            "□ 변경 관리 프로세스 정의",
            "□ 교육 및 훈련 계획 수립",
            "□ 롤백 계획 준비",
            "□ 프로젝트 마일스톤 및 일정 관리",
            "",
            "주요 위험:",
            "- 조직 내 저항 및 변화 거부",
            "- 이해관계자 간 의사소통 부재",
            "- 마이그레이션 일정 지연"
        ]),
        ("6. 벤더 및 파트너 위험", [
            "□ 클라우드 제공업체 SLA 검토",
            "□ 벤더 종속성(Lock-in) 평가",
            "□ 대체 제공업체 옵션 검토",
            "□ 계약 조건 및 책임 범위 명확화",
            "□ 기술 지원 및 에스컬레이션 프로세스 확인",
            "",
            "주요 위험:",
            "- 벤더의 서비스 품질 저하",
            "- 특정 벤더에 대한 과도한 의존",
            "- 계약 분쟁 또는 서비스 중단"
        ]),
        ("7. 위험 완화 및 대응 계획", [
            "각 위험 항목에 대한 완화 전략:",
            "- 고위험: 즉시 조치 및 지속적 모니터링",
            "- 중위험: 단계별 완화 계획 수립",
            "- 저위험: 정기 검토 및 예방 조치",
            "",
            "위험 발생 시 대응 절차:",
            "1. 위험 감지 및 보고",
            "2. 영향도 평가",
            "3. 대응팀 구성 및 역할 분담",
            "4. 완화 조치 실행",
            "5. 결과 모니터링 및 문서화",
            "6. 재발 방지 대책 수립"
        ])
    ]

    for section_title, items in sections:
        doc.add_heading(section_title, level=1)
        for item in items:
            if item.startswith("□"):
                doc.add_paragraph(item, style='List Bullet')
            elif item == "":
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph(item)

    # 결론
    doc.add_heading("8. 결론", level=1)
    conclusion = [
        "본 위험 관리 체크리스트는 클라우드 마이그레이션 프로젝트의 성공적인 수행을 위해 필수적으로 검토해야 할 항목들을 포함하고 있습니다.",
        "",
        "정기적으로 체크리스트를 검토하고 업데이트하여 새로운 위험을 조기에 식별하고 대응할 수 있도록 해야 합니다.",
        "",
        "모든 팀원이 위험 관리의 중요성을 인식하고 적극적으로 참여할 때 프로젝트의 성공 가능성이 높아집니다."
    ]

    for item in conclusion:
        doc.add_paragraph(item)

    docx_filename = "test_risk_checklist.docx"
    doc.save(docx_filename)
    print(f"DOCX 파일 생성 완료: {docx_filename}")

if __name__ == "__main__":
    create_test_pdf()
    create_test_docx()
    print("\n테스트 파일 생성이 완료되었습니다!")
    print("- test_risk_management.pdf")
    print("- test_risk_checklist.docx")
